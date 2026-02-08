"""Document Extraction Utility - Reusable module for extracting text from various file formats.

Supports:
- PDF files (local and from Google Drive)
- Google Docs (via Drive API export)
- Plain text files
- DOCX files (if python-docx installed)

This module provides a centralized way to extract text content from documents
for ingestion into RAG systems or other text processing pipelines.
"""
import io
import logging
from pathlib import Path
from typing import Optional, Union
from dataclasses import dataclass

logger = logging.getLogger(__name__)


@dataclass
class ExtractionResult:
    """Result of document extraction."""
    success: bool
    content: str
    source: str
    format: str
    page_count: int = 0
    error: Optional[str] = None
    metadata: dict = None
    
    def __post_init__(self):
        if self.metadata is None:
            self.metadata = {}


class DocumentExtractor:
    """Unified document text extraction utility.
    
    Example:
        >>> extractor = DocumentExtractor()
        >>> result = extractor.extract_from_file("document.pdf")
        >>> if result.success:
        ...     print(result.content)
        
        >>> # Extract from Google Drive
        >>> result = extractor.extract_from_drive(file_id, mime_type)
    """
    
    def __init__(self):
        """Initialize the extractor with available backends."""
        self._pypdf_available = self._check_pypdf()
        self._docx_available = self._check_docx()
        self._drive_provider = None
    
    def _check_pypdf(self) -> bool:
        """Check if pypdf is available."""
        try:
            import pypdf
            return True
        except ImportError:
            logger.warning("pypdf not installed. PDF extraction unavailable.")
            return False
    
    def _check_docx(self) -> bool:
        """Check if python-docx is available."""
        try:
            import docx
            return True
        except ImportError:
            logger.debug("python-docx not installed. DOCX extraction unavailable.")
            return False
    
    @property
    def drive_provider(self):
        """Lazy-load Google Drive provider."""
        if self._drive_provider is None:
            from .providers.google_drive_provider import GoogleDriveProvider
            self._drive_provider = GoogleDriveProvider()
        return self._drive_provider
    
    def extract_from_file(self, file_path: Union[str, Path]) -> ExtractionResult:
        """Extract text from a local file.
        
        Args:
            file_path: Path to the file to extract from.
            
        Returns:
            ExtractionResult with extracted content or error.
        """
        path = Path(file_path)
        
        if not path.exists():
            return ExtractionResult(
                success=False,
                content="",
                source=str(path),
                format="unknown",
                error=f"File not found: {path}"
            )
        
        suffix = path.suffix.lower()
        
        if suffix == '.pdf':
            return self._extract_pdf(path)
        elif suffix == '.docx':
            return self._extract_docx(path)
        elif suffix in ['.txt', '.md', '.json', '.csv']:
            return self._extract_text(path)
        else:
            return ExtractionResult(
                success=False,
                content="",
                source=str(path),
                format=suffix,
                error=f"Unsupported format: {suffix}"
            )
    
    def extract_from_bytes(
        self, 
        data: bytes, 
        format: str,
        source_name: str = "memory"
    ) -> ExtractionResult:
        """Extract text from in-memory bytes.
        
        Args:
            data: Raw bytes of the document.
            format: File format (e.g., 'pdf', 'docx').
            source_name: Name for logging/tracking.
            
        Returns:
            ExtractionResult with extracted content.
        """
        if format.lower() == 'pdf':
            return self._extract_pdf_bytes(data, source_name)
        elif format.lower() == 'docx':
            return self._extract_docx_bytes(data, source_name)
        else:
            return ExtractionResult(
                success=False,
                content="",
                source=source_name,
                format=format,
                error=f"Unsupported format: {format}"
            )
    
    def extract_from_drive(
        self, 
        file_id: str, 
        mime_type: str,
        file_name: str = ""
    ) -> ExtractionResult:
        """Extract text from a Google Drive file.
        
        Handles:
        - Google Docs (exported as text)
        - PDFs (downloaded and extracted)
        - Other text-exportable formats
        
        Args:
            file_id: Google Drive file ID.
            mime_type: MIME type of the file.
            file_name: Optional file name for metadata.
            
        Returns:
            ExtractionResult with extracted content.
        """
        from googleapiclient.http import MediaIoBaseDownload
        
        source = file_name or file_id
        
        try:
            # Google Docs - export as plain text
            if mime_type == 'application/vnd.google-apps.document':
                content = self.drive_provider.download_file(file_id, mime_type)
                return ExtractionResult(
                    success=True,
                    content=content,
                    source=source,
                    format="google_doc",
                    metadata={"mime_type": mime_type}
                )
            
            # Google Sheets - export as CSV
            elif mime_type == 'application/vnd.google-apps.spreadsheet':
                request = self.drive_provider.service.files().export_media(
                    fileId=file_id, 
                    mimeType='text/csv'
                )
                fh = io.BytesIO()
                downloader = MediaIoBaseDownload(fh, request)
                done = False
                while not done:
                    _, done = downloader.next_chunk()
                fh.seek(0)
                content = fh.read().decode('utf-8')
                return ExtractionResult(
                    success=True,
                    content=content,
                    source=source,
                    format="google_sheet",
                    metadata={"mime_type": mime_type}
                )
            
            # PDF - download and extract
            elif mime_type == 'application/pdf':
                request = self.drive_provider.service.files().get_media(fileId=file_id)
                fh = io.BytesIO()
                downloader = MediaIoBaseDownload(fh, request)
                done = False
                while not done:
                    _, done = downloader.next_chunk()
                fh.seek(0)
                return self._extract_pdf_bytes(fh.read(), source)
            
            # DOCX and other Office formats
            elif mime_type in [
                'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                'application/msword'
            ]:
                request = self.drive_provider.service.files().get_media(fileId=file_id)
                fh = io.BytesIO()
                downloader = MediaIoBaseDownload(fh, request)
                done = False
                while not done:
                    _, done = downloader.next_chunk()
                fh.seek(0)
                return self._extract_docx_bytes(fh.read(), source)
            
            # Plain text
            elif mime_type.startswith('text/'):
                request = self.drive_provider.service.files().get_media(fileId=file_id)
                fh = io.BytesIO()
                downloader = MediaIoBaseDownload(fh, request)
                done = False
                while not done:
                    _, done = downloader.next_chunk()
                fh.seek(0)
                content = fh.read().decode('utf-8')
                return ExtractionResult(
                    success=True,
                    content=content,
                    source=source,
                    format="text",
                    metadata={"mime_type": mime_type}
                )
            
            else:
                return ExtractionResult(
                    success=False,
                    content="",
                    source=source,
                    format=mime_type,
                    error=f"Unsupported MIME type: {mime_type}"
                )
                
        except Exception as e:
            logger.error(f"Drive extraction error for {source}: {e}")
            return ExtractionResult(
                success=False,
                content="",
                source=source,
                format=mime_type,
                error=str(e)
            )
    
    def _extract_pdf(self, path: Path) -> ExtractionResult:
        """Extract text from a local PDF file."""
        if not self._pypdf_available:
            return ExtractionResult(
                success=False,
                content="",
                source=str(path),
                format="pdf",
                error="pypdf not installed. Run: pip install pypdf"
            )
        
        try:
            import pypdf
            
            with open(path, 'rb') as f:
                reader = pypdf.PdfReader(f)
                pages = []
                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        pages.append(text)
                
                content = "\n\n".join(pages)
                
                return ExtractionResult(
                    success=True,
                    content=content,
                    source=str(path),
                    format="pdf",
                    page_count=len(reader.pages),
                    metadata={"filename": path.name}
                )
                
        except Exception as e:
            logger.error(f"PDF extraction error for {path}: {e}")
            return ExtractionResult(
                success=False,
                content="",
                source=str(path),
                format="pdf",
                error=str(e)
            )
    
    def _extract_pdf_bytes(self, data: bytes, source: str) -> ExtractionResult:
        """Extract text from PDF bytes."""
        if not self._pypdf_available:
            return ExtractionResult(
                success=False,
                content="",
                source=source,
                format="pdf",
                error="pypdf not installed"
            )
        
        try:
            import pypdf
            
            reader = pypdf.PdfReader(io.BytesIO(data))
            pages = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    pages.append(text)
            
            content = "\n\n".join(pages)
            
            return ExtractionResult(
                success=True,
                content=content,
                source=source,
                format="pdf",
                page_count=len(reader.pages)
            )
            
        except Exception as e:
            logger.error(f"PDF bytes extraction error: {e}")
            return ExtractionResult(
                success=False,
                content="",
                source=source,
                format="pdf",
                error=str(e)
            )
    
    def _extract_docx(self, path: Path) -> ExtractionResult:
        """Extract text from a local DOCX file."""
        if not self._docx_available:
            return ExtractionResult(
                success=False,
                content="",
                source=str(path),
                format="docx",
                error="python-docx not installed"
            )
        
        try:
            from docx import Document
            
            doc = Document(str(path))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            content = "\n".join(paragraphs)
            
            return ExtractionResult(
                success=True,
                content=content,
                source=str(path),
                format="docx",
                metadata={"filename": path.name}
            )
            
        except Exception as e:
            logger.error(f"DOCX extraction error for {path}: {e}")
            return ExtractionResult(
                success=False,
                content="",
                source=str(path),
                format="docx",
                error=str(e)
            )
    
    def _extract_docx_bytes(self, data: bytes, source: str) -> ExtractionResult:
        """Extract text from DOCX bytes."""
        if not self._docx_available:
            return ExtractionResult(
                success=False,
                content="",
                source=source,
                format="docx",
                error="python-docx not installed"
            )
        
        try:
            from docx import Document
            
            doc = Document(io.BytesIO(data))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            content = "\n".join(paragraphs)
            
            return ExtractionResult(
                success=True,
                content=content,
                source=source,
                format="docx"
            )
            
        except Exception as e:
            logger.error(f"DOCX bytes extraction error: {e}")
            return ExtractionResult(
                success=False,
                content="",
                source=source,
                format="docx",
                error=str(e)
            )
    
    def _extract_text(self, path: Path) -> ExtractionResult:
        """Extract content from a plain text file."""
        try:
            content = path.read_text(encoding='utf-8')
            return ExtractionResult(
                success=True,
                content=content,
                source=str(path),
                format="text",
                metadata={"filename": path.name}
            )
        except Exception as e:
            return ExtractionResult(
                success=False,
                content="",
                source=str(path),
                format="text",
                error=str(e)
            )


# Convenience function for quick extraction
def extract_text(source: Union[str, Path, bytes], format: str = None) -> str:
    """Quick text extraction from various sources.
    
    Args:
        source: File path or bytes.
        format: Required if source is bytes.
        
    Returns:
        Extracted text content, or empty string on failure.
    """
    extractor = DocumentExtractor()
    
    if isinstance(source, bytes):
        if not format:
            raise ValueError("format is required when source is bytes")
        result = extractor.extract_from_bytes(source, format)
    else:
        result = extractor.extract_from_file(source)
    
    return result.content if result.success else ""
